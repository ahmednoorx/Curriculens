import streamlit as st
import os
import re
from src.models.text_generation import TextGenerator
from src.parsers.pdf_parser import PDFParser
from src.parsers.docx_parser import DOCXParser
from io import BytesIO
from fpdf import FPDF
from docx import Document
from streamlit_chat import message
from dotenv import load_dotenv
import nltk
from collections import Counter
from nltk.corpus import stopwords

# Download NLTK stopwords if not already present
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

load_dotenv()  # This loads environment variables from .env file
api_key = os.getenv("OPENAI_API_KEY")

# Improved chapter extraction regex
CHAPTER_PATTERNS = [
    r"(?:Chapter\s+(\d+)[\s:\-]+)([^\n]+)",
    r"(?:CHAPTER\s+(\d+)[\s:\-]+)([^\n]+)",
    r"(?:CHAPTER\s+(\d+))",
    r"(?:Section\s+(\d+)[\s:\-]+)([^\n]+)",
]

def extract_chapters(text):
    chapters = {}
    for pattern in CHAPTER_PATTERNS:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            if isinstance(match, tuple) and len(match) == 2:
                num, name = match
                chapters[int(num)] = name.strip()
            elif isinstance(match, str):
                chapters[int(match)] = f"Chapter {match}"
    return chapters

# Keyword extraction using NLTK
@st.cache_data
def extract_keywords(text, num_keywords=10):
    words = re.findall(r'\b\w+\b', text.lower())
    stop_words = set(stopwords.words('english'))
    filtered = [w for w in words if w not in stop_words and len(w) > 2]
    most_common = Counter(filtered).most_common(num_keywords)
    return [w for w, _ in most_common]


def export_to_pdf(text, filename="output.pdf"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    for line in text.split('\n'):
        pdf.multi_cell(0, 10, line)
    pdf_output = pdf.output(dest='S').encode('latin1')
    pdf_buffer = BytesIO(pdf_output)
    pdf_buffer.seek(0)
    return pdf_buffer


def export_to_word(text, filename="output.docx"):
    doc = Document()
    doc.add_paragraph(text)
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer


def chatbot_response(user_input, extracted_text, chapters, content_type, selected_chapter=None):
    context = ""

    if content_type == "Book" and selected_chapter and selected_chapter in chapters:
        chapter_title = f"Chapter {selected_chapter}: {chapters[selected_chapter]}"
        pattern = re.compile(
            rf"{re.escape(chapter_title)}(.*?)(Chapter\s+\d+:|$)", re.DOTALL | re.IGNORECASE
        )
        match = pattern.search(extracted_text)
        if match:
            context = match.group(1).strip()
        else:
            context = extracted_text
    else:
        context = extracted_text

    context = context[:3000]

    prompt = f"""You are an educational assistant. The user uploaded the following content:

{context}

User question: {user_input}

Please provide a clear, concise and helpful answer based on the above content."""

    text_generator = TextGenerator()
    return text_generator.generate_text(prompt)


def main():
    st.set_page_config(layout="wide")
    st.title("📚 Curriculum Assistant")

    # Theme toggle
    theme = st.sidebar.radio("Theme", ["Light", "Dark"], index=0)
    if theme == "Dark":
        st.markdown(
            """
            <style>
            body, .stApp { background-color: #222 !important; color: #eee !important; }
            </style>
            """,
            unsafe_allow_html=True,
        )

    # Sidebar organization with expanders
    with st.sidebar:
        with st.expander("Instructions", expanded=True):
            st.markdown("""
            **Instructions:**
            1. Upload a PDF or DOCX file.
            2. Select content type (Syllabus/Book).
            3. Explore chapters, generate content, or chat with Curriculens.
            4. Export your results.
            """)
        with st.expander("Settings", expanded=True):
            st.radio("Theme", ["Light", "Dark"], index=0, key="theme_radio")
        with st.expander("Reset", expanded=False):
            if st.button("Reset Session"):
                for key in ["extracted_text", "chapters", "chat_history", "all_generated"]:
                    if key in st.session_state:
                        del st.session_state[key]
                st.rerun()

    if "extracted_text" not in st.session_state:
        st.session_state.extracted_text = None
    if "chapters" not in st.session_state:
        st.session_state.chapters = {}
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "all_generated" not in st.session_state:
        st.session_state.all_generated = []

    content_type = st.radio("What are you uploading?", ["Syllabus", "Book"])

    uploaded_file = st.file_uploader("Upload a PDF or DOCX file", type=["pdf", "docx"])

    # Limit input size for uploaded files
    MAX_CHARS = 12000  # adjust as needed for model/Spaces
    if uploaded_file is not None:
        with st.spinner("Extracting text..."):
            try:
                if uploaded_file.name.lower().endswith(".pdf"):
                    parser = PDFParser()
                elif uploaded_file.name.lower().endswith(".docx"):
                    parser = DOCXParser()
                else:
                    st.error("Unsupported file type.")
                    return
                extracted_text = parser.extract_text(uploaded_file)
                if len(extracted_text) > MAX_CHARS:
                    st.warning(f"Uploaded document is large. Only the first {MAX_CHARS} characters will be used.")
                    extracted_text = extracted_text[:MAX_CHARS]
                st.session_state.extracted_text = extracted_text
                st.success("✅ File uploaded and text extracted!")
            except Exception as e:
                st.error(f"Error extracting text: {e}")
                return

        if content_type == "Book":
            chapters = extract_chapters(extracted_text)
            if chapters:
                st.session_state.chapters = chapters
                st.subheader("📖 Detected Chapters")
                for num, name in chapters.items():
                    st.markdown(f"- **Chapter {num}:** {name}")
            else:
                st.warning("⚠️ No chapters detected. You can mark chapters manually below.")
                # Optionally, allow manual chapter marking (not implemented for brevity)

    if st.session_state.extracted_text:
        col1, col2 = st.columns([2, 1])
        with col1:
            st.header("📌 Generate Content")
            task = st.selectbox("Select task", ["Lesson Plan", "MCQs", "Short Questions", "Summarize", "Extract Keywords"])
            export_format = st.selectbox("Select export format", ["PDF", "Word"])
            selected_chapter = None
            chapter_text = st.session_state.extracted_text
            if content_type == "Book" and st.session_state.chapters:
                chapter_options = [f"Chapter {num}: {name}" for num, name in st.session_state.chapters.items()]
                chapter_choice = st.selectbox("Select chapter (optional)", ["All"] + chapter_options)
                if chapter_choice != "All":
                    selected_chapter = int(chapter_choice.split()[1].replace(":", ""))
                    # Extract only the selected chapter's text
                    chapter_title = f"Chapter {selected_chapter}: {st.session_state.chapters[selected_chapter]}"
                    pattern = re.compile(rf"{re.escape(chapter_title)}(.*?)(Chapter\\s+\\d+:|$)", re.DOTALL | re.IGNORECASE)
                    match = pattern.search(st.session_state.extracted_text)
                    if match:
                        chapter_text = match.group(1).strip()
                    else:
                        chapter_text = st.session_state.extracted_text
            if st.button("Generate"):
                progress = st.progress(0, text="Generating content...")
                try:
                    with st.spinner("Generating content..."):
                        progress.progress(10)
                        if task == "Extract Keywords":
                            keywords = extract_keywords(chapter_text)
                            result = "Keywords: " + ", ".join(keywords)
                        elif task == "Summarize":
                            prompt = f"Summarize the following content:\n\n{chapter_text[:3000]}"
                            tg = TextGenerator()
                            progress.progress(30)
                            result = tg.generate_text(prompt)
                        else:
                            prompt_map = {
                                "Lesson Plan": "Create a lesson plan for the following content:",
                                "MCQs": "Generate multiple choice questions for the following content:",
                                "Short Questions": "Generate short answer questions for the following content:"
                            }
                            prompt = f"{prompt_map[task]}\n\n{chapter_text[:3000]}"
                            tg = TextGenerator()
                            progress.progress(30)
                            result = tg.generate_text(prompt)
                        progress.progress(90)
                        st.text_area("Generated Content", value=result, height=200)
                        st.session_state.all_generated.append(result)
                        progress.progress(100)
                except Exception as e:
                    st.error(f"Error during content generation: {e}")
                if st.button("Export All Generated Content"):
                    all_text = "\n\n".join(st.session_state.all_generated)
                    if export_format == "PDF":
                        buf = export_to_pdf(all_text)
                        st.download_button("Download PDF", buf, file_name="all_content.pdf")
                    else:
                        buf = export_to_word(all_text)
                        st.download_button("Download Word", buf, file_name="all_content.docx")
        with col2:
            st.header("🤖 Curriculens")
            selected_chapter = None
            if content_type == "Book" and st.session_state.chapters:
                chapter_options = [f"Chapter {num}: {name}" for num, name in st.session_state.chapters.items()]
                chapter_choice = st.selectbox("Chat about chapter (optional)", ["All"] + chapter_options, key="chat_chapter")
                if chapter_choice != "All":
                    selected_chapter = int(chapter_choice.split()[1].replace(":", ""))
            user_input = st.text_input("You:", key="chat_input")
            if user_input:
                progress = st.progress(0, text="Generating response...")
                try:
                    with st.spinner("Generating response..."):
                        progress.progress(10)
                        response = chatbot_response(user_input, st.session_state.extracted_text, st.session_state.chapters, content_type, selected_chapter)
                        progress.progress(90)
                        st.session_state.chat_history.append({"user": user_input, "bot": response, "chapter": selected_chapter})
                        progress.progress(100)
                except Exception as e:
                    st.error(f"Error during chat response: {e}")
            # Clean, ordered chat history display (most recent at top)
            st.markdown("#### Chat History")
            for i, entry in enumerate(reversed(st.session_state.chat_history)):
                chapter_info = f" (Chapter {entry['chapter']})" if entry['chapter'] else ""
                st.markdown(f"**You{chapter_info}:** {entry['user']}")
                st.markdown(f"**Curriculens:** {entry['bot']}")
                st.markdown("---")
            # Export chat history
            def chat_history_to_text():
                lines = []
                for entry in reversed(st.session_state.chat_history):
                    chapter_info = f" (Chapter {entry['chapter']})" if entry['chapter'] else ""
                    lines.append(f"You{chapter_info}: {entry['user']}")
                    lines.append(f"Curriculens: {entry['bot']}")
                    lines.append("")
                return "\n".join(lines)
            def chat_history_to_word():
                doc = Document()
                for entry in reversed(st.session_state.chat_history):
                    chapter_info = f" (Chapter {entry['chapter']})" if entry['chapter'] else ""
                    doc.add_paragraph(f"You{chapter_info}: {entry['user']}")
                    doc.add_paragraph(f"Curriculens: {entry['bot']}")
                    doc.add_paragraph("")
                buf = BytesIO()
                doc.save(buf)
                buf.seek(0)
                return buf
            st.markdown("**Export Chat History:**")
            st.download_button("Download as TXT", chat_history_to_text(), file_name="chat_history.txt")
            st.download_button("Download as Word", chat_history_to_word(), file_name="chat_history.docx")

if __name__ == "__main__":
    main()
